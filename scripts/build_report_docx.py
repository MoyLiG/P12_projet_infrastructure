"""
Génère le rapport projet `.docx` à partir du contenu structuré ci-dessous.

Usage :
    python scripts/build_report_docx.py
    -> écrit docs/Le_Gall_Morgan_Option_B_rapport_062026.docx

Le `.docx` reste **synthétique** (10-15 pages). Le détail long (chronologie,
alternatives écartées, justifications fines) vit dans `journal.md`.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "Le_Gall_Morgan_Option_B_rapport_062026.docx"


# ----- helpers -------------------------------------------------------

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def add_pagebreak(doc):
    doc.add_page_break()


# ----- contenu --------------------------------------------------------

def build_document() -> Document:
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # --- Page de garde -------------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("\n\n\nPOC Avantages Sportifs")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Sport Data Solution — Pipeline ETL bout-en-bout\n")
    sr.font.size = Pt(14)
    sr.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        "\n\n\n\n\n\n\nProjet OpenClassrooms P12 — Data Engineer\n"
        "Option B\n\n"
        "Morgan Le Gall\n"
        "Juin 2026"
    )
    mr.font.size = Pt(13)

    add_pagebreak(doc)

    # --- 1. Contexte ----------------------------------------------------
    add_heading(doc, "1. Contexte et objectifs", 1)
    add_para(doc,
        "Sport Data Solution est une start-up qui développe des solutions de monitoring "
        "sportif. Juliette, cofondatrice, souhaite tester un dispositif d'avantages "
        "pour ses 162 salariés afin d'encourager la pratique sportive :")
    add_bullets(doc, [
        "Prime sportive : +5 % du salaire annuel brut pour les salariés venant au "
        "bureau en mode actif (marche, course, vélo, trottinette).",
        "5 journées bien-être par an pour les salariés ayant ≥ 15 activités "
        "physiques déclarées sur 12 mois.",
    ])
    add_para(doc,
        "Le POC livre un pipeline ETL bout-en-bout qui prouve la faisabilité technique, "
        "garantit la qualité des données, et calcule l'impact financier des deux "
        "avantages. La consigne exige une démo live (modifier le taux de prime, insérer "
        "une activité en direct) et une restitution PowerBI.")

    add_heading(doc, "Livrables", 2)
    add_bullets(doc, [
        "Code source sur GitHub avec README détaillé.",
        "Pipeline orchestré (Kestra) avec monitoring et alertes.",
        "Dashboard PowerBI consommant la base de données.",
        "Support de présentation pour la soutenance.",
    ])

    add_pagebreak(doc)

    # --- 2. Architecture -----------------------------------------------
    add_heading(doc, "2. Architecture technique", 1)
    add_para(doc, "Stack retenue après comparaison :")
    add_table(doc,
        headers=["Brique", "Choix", "Justification"],
        rows=[
            ["Orchestration", "Kestra 0.19",
             "Flow YAML versionné, replay natif (exigence consigne), UI de monitoring."],
            ["Base de données", "PostgreSQL 16",
             "Rôles, Row Level Security, triggers d'audit — adapté aux PII RH."],
            ["Langage", "Python 3.11",
             "pandas / SQLAlchemy / Faker / Great Expectations."],
            ["Tests qualité", "Great Expectations 1.x",
             "Cité dans la note de cadrage. Suite déclarative, rapport JSON."],
            ["Géocodage", "Google Maps Distance Matrix",
             "+ cache PG + fallback haversine. 162 appels par run = quasi gratuit."],
            ["Notification", "Slack Incoming Webhook",
             "+ fallback fichier si le webhook répond en erreur."],
            ["Restitution", "PowerBI Desktop",
             "Connecteur PG natif, paramètres dynamiques pour la démo live."],
            ["Conteneurisation", "Docker Compose",
             "Une commande pour reconstruire la stack from scratch."],
        ])

    add_heading(doc, "Schémas PostgreSQL (pattern dbt-like)", 2)
    add_bullets(doc, [
        "raw — copie 1:1 des XLSX. Jamais modifiée (audit trail).",
        "staging — typage strict + contraintes CHECK + clés étrangères.",
        "marts — tables consommées par PowerBI : eligibility_prime, eligibility_wellbeing, kpi_by_bu.",
        "audit — run_log (volumétrie/durée par étape) et eligibility_changes (traçabilité RGPD).",
        "cache — gmaps_distance pour idempotence + maîtrise du coût API.",
    ])

    add_heading(doc, "Flux du pipeline", 2)
    add_para(doc,
        "Le flow Kestra enchaîne sept étapes : extraction RH, extraction sport, "
        "génération des activités (Strava-like), validation géographique, validation "
        "qualité Great Expectations, calcul des avantages, notification Slack. Une "
        "défaillance à n'importe quelle étape interrompt le flow et déclenche une alerte mail.")

    add_pagebreak(doc)

    # --- 3. Données -----------------------------------------------------
    add_heading(doc, "3. Données", 1)
    add_heading(doc, "Sources", 2)
    add_bullets(doc, [
        "Données+RH.xlsx (162 lignes × 11 colonnes) — ID, nom, prénom, date naissance, "
        "BU, date embauche, salaire brut, contrat, CP, adresse, moyen de déplacement.",
        "Données+Sportive.xlsx (1000 lignes × 2 colonnes) — ID salarié, sport déclaré "
        "(nombreux NULL filtrés à l'extraction).",
        "Activités sportives — générées par src/generate/activities.py (Faker + numpy) "
        "sur les 12 derniers mois, environ 4 000 lignes, cohérentes avec le sport "
        "déclaré de chaque salarié.",
    ])

    add_heading(doc, "Données sensibles (PII)", 2)
    add_para(doc, "Le fichier RH contient des PII (nom, salaire, adresse). Le projet applique :")
    add_bullets(doc, [
        "Aucun fichier XLSX versionné — .gitignore strict.",
        "Trois rôles PG distincts : etl_writer (R/W), analyst_reader / powerbi_reader "
        "(read marts only, jamais d'accès à staging.employees).",
        "Pseudonymisation : SHA-256 salé du couple nom+prénom pour les analystes.",
        "Row Level Security activée sur staging.employees.",
        "Triggers d'audit sur les marts → trace exhaustive des calculs.",
        "Inserts paramétrés (SQLAlchemy) — aucune concaténation SQL.",
    ])

    add_pagebreak(doc)

    # --- 4. Qualité des données ----------------------------------------
    add_heading(doc, "4. Qualité des données", 1)
    add_para(doc,
        "Trois lignes de défense complémentaires garantissent l'intégrité :")

    add_heading(doc, "Ligne 1 — Contraintes SQL (au chargement)", 2)
    add_bullets(doc, [
        "salaire_brut > 0, jours_cp >= 0.",
        "moyen_deplacement IN (enum fermé) — valeur inconnue → erreur immédiate.",
        "end_dt > start_dt, distance_m IS NULL OR distance_m >= 0.",
        "Clés étrangères vérifiées (id_employee de staging.activities doit exister).",
    ])

    add_heading(doc, "Ligne 2 — Great Expectations (avant calculs)", 2)
    add_para(doc,
        "Suite déclarative sur employees et activities. Le rapport JSON est écrit "
        "dans data/ge_docs/ avec un timestamp par run. Tout échec lève une exception "
        "qui interrompt le pipeline Kestra et déclenche l'alerte mail.")
    add_bullets(doc, [
        "Unicité de id_employee.",
        "salaire_brut ∈ [10k, 300k €].",
        "moyen_deplacement ∈ enum.",
        "Dates d'activités non futures, distance ∈ [0, 200 km].",
    ])

    add_heading(doc, "Ligne 3 — Validation métier géographique", 2)
    add_para(doc,
        "Pour les salariés déclarant venir en mode actif, Google Maps Distance Matrix "
        "calcule la distance domicile → entreprise. Au-delà d'un seuil (15 km marche, "
        "25 km vélo), la déclaration est flaguée is_declaration_suspect = TRUE → "
        "le salarié n'est pas éligible à la prime, et l'anomalie remonte au dashboard.")

    add_pagebreak(doc)

    # --- 5. Calculs métier ---------------------------------------------
    add_heading(doc, "5. Calculs métier", 1)
    add_heading(doc, "Prime sportive", 2)
    add_para(doc, "Formule (SQL paramétré, UPSERT idempotent) :", italic=True)
    add_para(doc,
        "prime_amount = prime_rate × salaire_brut  SI  "
        "moyen_deplacement ∈ {Marche/running, Vélo/Trottinette}  "
        "ET  is_declaration_suspect = FALSE")

    add_heading(doc, "Jours bien-être", 2)
    add_para(doc,
        "5 jours accordés au salarié si COUNT(activités sur 365 derniers jours) ≥ "
        "wellbeing_threshold (15 par défaut).")

    add_heading(doc, "Paramètres ajustables sans redéploiement", 2)
    add_para(doc,
        "Les deux paramètres clés (prime_rate, wellbeing_threshold) sont des inputs "
        "Kestra — modifiables au moment du Execute dans l'UI ou via l'API. "
        "C'est cette propriété qui rend la démo live possible.")

    add_pagebreak(doc)

    # --- 6. Monitoring -------------------------------------------------
    add_heading(doc, "6. Monitoring et observabilité", 1)
    add_bullets(doc, [
        "Table audit.run_log : 1 ligne par étape avec rows_in, rows_out, "
        "duration_ms, status, message. Visualisable dans une page dédiée PowerBI.",
        "UI Kestra : graphe d'exécution, logs en temps réel, replay manuel.",
        "Alerte mail (errors:) déclenchée sur tout échec.",
        "Cache.gmaps_distance : aussi un signal de santé — un cache vide après "
        "plusieurs runs trahit un problème d'API key.",
    ])

    add_pagebreak(doc)

    # --- 7. Sécurité ---------------------------------------------------
    add_heading(doc, "7. Sécurité et RGPD", 1)
    add_bullets(doc, [
        "Données RH considérées sensibles dès la phase de design.",
        "Aucun secret versionné — .env git-ignored, .env.example en placeholders.",
        "Chiffrement at-rest : volume PG sur disque chiffré (BitLocker sous Windows).",
        "Pseudonymisation pour les analystes (SHA-256 salé).",
        "Audit trail RGPD via triggers + table audit.eligibility_changes (qui a "
        "calculé quoi, quand, avec quels paramètres).",
        "Script forget_employee.py prévu pour le droit à l'oubli (next step prod).",
        "Row Level Security pour interdire la lecture brute de staging.employees.",
    ])

    add_pagebreak(doc)

    # --- 8. Tests ------------------------------------------------------
    add_heading(doc, "8. Tests automatisés", 1)
    add_para(doc, "Tests pytest organisés en deux familles :")
    add_bullets(doc, [
        "Unitaires (rapides, sans DB) : normalisation transport, hash PII, "
        "haversine, formatage Slack, distributions du générateur.",
        "Intégration (RUN_INTEGRATION=1 + Postgres) : golden values des calculs "
        "métier, seuil bien-être à la frontière, changement de taux de prime.",
    ])
    add_para(doc, "Stratégie golden values : 3 cas câblés manuellement —")
    add_bullets(doc, [
        "Transport actif + déclaration cohérente → éligible (1 500 € pour 30 k brut).",
        "Transport actif + suspect → non éligible (raison explicite).",
        "Transport non actif → non éligible.",
    ])

    add_pagebreak(doc)

    # --- 9. Démo soutenance --------------------------------------------
    add_heading(doc, "9. Démo soutenance — scénarios", 1)
    add_para(doc, "Conformément à la consigne OC, deux scénarios sont prêts (scripts/demo.sh) :")
    add_heading(doc, "Scénario A — changement du taux de prime", 2)
    add_bullets(doc, [
        "Run du flow avec prime_rate = 0.05.",
        "Lecture des KPI dans PowerBI : coût total = X €.",
        "Re-run avec prime_rate = 0.07.",
        "Refresh PowerBI : coût total recalculé. La table marts.eligibility_prime "
        "garde les deux versions via le champ run_id pour comparaison.",
    ])
    add_heading(doc, "Scénario B — nouvelle activité en live", 2)
    add_bullets(doc, [
        "INSERT manuel dans staging.activities (script ou SQL direct).",
        "Re-run de l'étape post_slack uniquement.",
        "Vérification : message reçu dans #sport-data + ligne visible dans "
        "PowerBI (page Détail Bien-être).",
    ])

    add_pagebreak(doc)

    # --- 10. Limites + perspectives ------------------------------------
    add_heading(doc, "10. Limites et perspectives", 1)
    add_heading(doc, "Hors scope du POC", 2)
    add_bullets(doc, [
        "Pas d'intégration vraie Strava (générateur Faker + numpy à la place).",
        "Pas de CI/CD (GitHub Actions est la prochaine étape).",
        "Pas de scale-out (Kubernetes, Spark) — surdimensionné pour 162 employés.",
    ])
    add_heading(doc, "Perspectives", 2)
    add_bullets(doc, [
        "Connexion API Strava réelle (OAuth) en remplacement du générateur.",
        "TTL sur le cache Google Maps (rafraîchissement périodique).",
        "Image Docker custom avec dépendances pré-installées (gain ~30 s/étape).",
        "Dashboard temps réel via Kestra → Grafana plutôt que PowerBI mensuel.",
        "Politique de rotation du sel PII + ré-hash automatique.",
    ])

    return doc


if __name__ == "__main__":
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUT)
    print(f"OK — rapport écrit : {OUT}")
